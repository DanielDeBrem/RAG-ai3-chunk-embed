# document_loader.py
from __future__ import annotations

import os
from typing import Literal

from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook


SupportedExt = Literal[".pdf", ".docx", ".xlsx", ".txt", ".md"]


def detect_ext(path: str) -> str:
    return os.path.splitext(path)[1].lower()


def load_pdf(path: str) -> str:
    reader = PdfReader(path)
    texts = []
    for page in reader.pages:
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        if txt.strip():
            texts.append(txt)
    return "\n\n".join(texts)


def load_docx(path: str) -> str:
    doc = DocxDocument(path)
    parts = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)
    # optioneel: tabellen meenemen
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def load_xlsx(path: str) -> str:
    wb = load_workbook(path, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"### Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def load_text_like(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def load_file_as_text(path: str) -> str:
    if not os.path.isfile(path):
        raise FileNotFoundError(path)

    ext = detect_ext(path)

    if ext == ".pdf":
        return load_pdf(path)
    if ext == ".docx":
        return load_docx(path)
    if ext == ".xlsx":
        return load_xlsx(path)
    if ext in {".txt", ".md"}:
        return load_text_like(path)

    # fallback: probeer als text
    return load_text_like(path)
